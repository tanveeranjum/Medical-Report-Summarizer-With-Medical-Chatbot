# requirements.txt content:
"""
gradio==4.19.2
transformers==4.36.2
torch==2.1.0
torchvision==0.16.0
torchaudio==2.1.0
pdf2image==1.16.3
pytesseract==0.3.10
python-docx==1.1.0
PyPDF2==3.0.1
Pillow==10.1.0
easyocr==1.7.1
sentencepiece==0.1.99
accelerate==0.25.0
groq==0.3.0
numpy==1.24.3
"""

import os
import re
import tempfile
import traceback
from typing import Dict, List, Optional
import gradio as gr

# Optional imports with graceful fallback
try:
    # PDF processing
    import PyPDF2
    from pdf2image import convert_from_bytes
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    # Image processing
    import pytesseract
    from PIL import Image
    import easyocr
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    # Word document processing
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    # AI/ML
    from transformers import pipeline
    import torch
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False

try:
    # Groq for faster inference
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False

# ==================== CONFIGURATION ====================

# Get API key from environment variable (safe for Hugging Face)
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
ENABLE_GROQ = bool(GROQ_API_KEY)

# Model configurations
MEDICAL_NER_MODEL = "samrawal/bert-base-uncased_clinical-ner"  # Smaller, faster model
AVAILABLE_GROQ_MODELS = ["llama-3.1-8b-instant", "mixtral-8x7b-32768"]

# ==================== DOCUMENT PROCESSOR ====================

class MultiFormatDocumentProcessor:
    def __init__(self):
        self.reader = None
        if OCR_AVAILABLE:
            try:
                self.reader = easyocr.Reader(['en'])
            except:
                self.reader = None
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        return text.strip()
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from various file formats with fallbacks"""
        file_type = file_type.lower()
        
        # Text files
        if file_type == 'txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return self.clean_text(f.read())
            except:
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        return self.clean_text(f.read())
                except:
                    return "Error reading text file"
        
        # PDF files
        elif file_type == 'pdf':
            if not PDF_AVAILABLE:
                return "PDF processing not available"
            
            try:
                # Try PyPDF2 first
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if text.strip():
                    return self.clean_text(text)
                
                # Fallback to OCR if no text found
                if OCR_AVAILABLE:
                    images = convert_from_bytes(open(file_path, 'rb').read())
                    text = ""
                    for image in images:
                        text += pytesseract.image_to_string(image) + "\n"
                    return self.clean_text(text)
                
                return "Could not extract text from PDF"
                
            except Exception as e:
                return f"Error processing PDF: {str(e)}"
        
        # Image files
        elif file_type in ['jpg', 'jpeg', 'png', 'bmp']:
            if not OCR_AVAILABLE:
                return "OCR processing not available"
            
            try:
                text = ""
                if self.reader:
                    results = self.reader.readtext(file_path)
                    text = " ".join([result[1] for result in results])
                
                if not text.strip() and OCR_AVAILABLE:
                    image = Image.open(file_path)
                    text = pytesseract.image_to_string(image)
                
                return self.clean_text(text) if text.strip() else "No text found in image"
                
            except Exception as e:
                return f"Error processing image: {str(e)}"
        
        # Word documents
        elif file_type in ['docx', 'doc']:
            if not DOCX_AVAILABLE:
                return "Word document processing not available"
            
            try:
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                
                # Extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + " "
                        text += "\n"
                
                return self.clean_text(text)
                
            except Exception as e:
                return f"Error processing Word document: {str(e)}"
        
        else:
            return f"Unsupported file type: {file_type}"

# ==================== SERIOUSNESS ANALYZER ====================

class SeriousnessAnalyzer:
    def __init__(self):
        self.critical_terms = {
            "high": [
                "cancer", "malignant", "metastasis", "tumor", "carcinoma",
                "heart attack", "myocardial infarction", "stroke", "aneurysm",
                "sepsis", "organ failure", "critical condition", "emergency",
                "life-threatening", "rupture", "internal bleeding", "arrest"
            ],
            "medium": [
                "infection", "inflammation", "hypertension", "diabetes",
                "arthritis", "pneumonia", "bronchitis", "fracture",
                "ulcer", "kidney disease", "liver disease", "moderate",
                "worsening", "progressive", "chronic"
            ],
            "low": [
                "mild", "slight", "minor", "stable", "improving",
                "benign", "routine", "checkup", "follow-up", "observation"
            ]
        }
    
    def analyze_seriousness(self, text: str) -> Dict:
        """Analyze the overall seriousness of medical findings"""
        text_lower = text.lower()
        
        severity_scores = {"high": 0, "medium": 0, "low": 0}
        for severity, terms in self.critical_terms.items():
            for term in terms:
                if term in text_lower:
                    severity_scores[severity] += text_lower.count(term)
        
        # Calculate overall score
        overall_score = (
            severity_scores["high"] * 3 +
            severity_scores["medium"] * 2 +
            severity_scores["low"] * 1
        )
        
        # Determine level
        if overall_score >= 5 or severity_scores["high"] >= 2:
            seriousness_level = "High"
            recommendation = "🔴 URGENT: Consult healthcare provider immediately."
        elif overall_score >= 3:
            seriousness_level = "Medium"
            recommendation = "🟡 MODERATE: Schedule follow-up with your doctor."
        else:
            seriousness_level = "Low"
            recommendation = "🟢 ROUTINE: Discuss at your next appointment."
        
        return {
            "level": seriousness_level,
            "score": overall_score,
            "recommendation": recommendation,
            "term_counts": severity_scores
        }

# ==================== CONSULTANT SEARCH ====================

class PakistanConsultantSearch:
    def __init__(self):
        self.doctors_database = [
            {
                "id": 101,
                "name": "Dr. Ahmed Raza",
                "specialty": "Cardiology",
                "subspecialties": ["Hypertension", "Heart Disease", "Angioplasty"],
                "hospital": "Aga Khan University Hospital, Karachi",
                "address": "Stadium Road, Karachi, Pakistan",
                "phone": "+92-21-111-111-111",
                "rating": 4.8,
                "experience": 18,
                "accepting_patients": True,
                "online_consultation": True,
                "consultation_fee": "₨ 3,000",
                "languages": ["English", "Urdu"],
                "city": "Karachi"
            },
            {
                "id": 102,
                "name": "Dr. Saima Khan",
                "specialty": "Endocrinology",
                "subspecialties": ["Diabetes", "Thyroid", "Metabolic Disorders"],
                "hospital": "Shaukat Khanum Memorial Hospital, Lahore",
                "address": "H-Block, Johar Town, Lahore, Pakistan",
                "phone": "+92-42-111-111-111",
                "rating": 4.7,
                "experience": 15,
                "accepting_patients": True,
                "online_consultation": True,
                "consultation_fee": "₨ 2,500",
                "languages": ["English", "Urdu", "Punjabi"],
                "city": "Lahore"
            },
            {
                "id": 201,
                "name": "Dr. Online Consultant Team",
                "specialty": "General Physician",
                "subspecialties": ["Online Consultation", "Prescription", "Follow-up"],
                "hospital": "Telemedicine Pakistan",
                "address": "Online - Nationwide Service",
                "phone": "0300-123-4567",
                "rating": 4.4,
                "experience": 8,
                "accepting_patients": True,
                "online_consultation": True,
                "consultation_fee": "₨ 1,500",
                "languages": ["English", "Urdu"],
                "city": "Online"
            }
        ]
        
        self.condition_to_specialty = {
            "hypertension": "Cardiology",
            "diabetes": "Endocrinology",
            "stroke": "Neurology",
            "heart disease": "Cardiology",
            "thyroid": "Endocrinology",
            "fracture": "Orthopedics",
            "skin rash": "Dermatology",
            "pregnancy": "Gynecology",
            "child health": "Pediatrics"
        }
    
    def suggest_specialties(self, medical_terms: List[str]) -> List[str]:
        """Suggest specialties based on medical terms"""
        specialties = set()
        
        for term in medical_terms:
            term_lower = term.lower()
            for condition, specialty in self.condition_to_specialty.items():
                if condition in term_lower:
                    specialties.add(specialty)
        
        if not specialties:
            specialties = {"General Physician", "Family Medicine"}
        
        return list(specialties)[:3]  # Limit to 3 suggestions
    
    def search_consultants(self, specialty: str, city: str = "Any", online_only: bool = False) -> List[Dict]:
        """Search for consultants"""
        results = []
        
        for doctor in self.doctors_database:
            # Specialty filter
            specialty_match = (
                specialty.lower() in doctor["specialty"].lower() or
                any(specialty.lower() in sub.lower() for sub in doctor["subspecialties"])
            )
            
            # Location filter
            location_match = False
            if online_only:
                location_match = doctor["online_consultation"]
            elif city.lower() == "any":
                location_match = True
            else:
                location_match = city.lower() in doctor["city"].lower()
            
            if specialty_match and location_match:
                results.append(doctor)
        
        # Sort by rating
        results.sort(key=lambda x: x["rating"], reverse=True)
        return results[:5]  # Limit to 5 results
    
    def format_doctor_card(self, doctor: Dict) -> str:
        """Format doctor info for display"""
        return f"""
**👨‍⚕️ {doctor['name']}** ({doctor['rating']}⭐)
**Specialty**: {doctor['specialty']}
**Hospital**: {doctor['hospital']}
**Location**: {doctor['city']}
**Fee**: {doctor['consultation_fee']}
**Phone**: {doctor['phone']}
**Online**: {'✅ Yes' if doctor['online_consultation'] else '❌ No'}
---
"""

# ==================== AI MEDICAL ASSISTANT ====================

class EnhancedMedicalAIAssistant:
    def __init__(self, groq_api_key: str = ""):
        self.groq_api_key = groq_api_key
        self.seriousness_analyzer = SeriousnessAnalyzer()
        self.medical_ner = None
        self.groq_client = None
        
        # Medical glossary
        self.medical_glossary = {
            "hypertension": "high blood pressure",
            "myocardial infarction": "heart attack",
            "diabetes": "high blood sugar condition",
            "fracture": "broken bone",
            "benign": "not cancerous",
            "malignant": "cancerous"
        }
        
        # Initialize NER model
        if TRANSFORMERS_AVAILABLE:
            try:
                self.medical_ner = pipeline(
                    "ner",
                    model=MEDICAL_NER_MODEL,
                    aggregation_strategy="simple"
                )
            except Exception as e:
                print(f"Warning: Could not load NER model: {e}")
        
        # Initialize Groq client
        if GROQ_AVAILABLE and groq_api_key:
            try:
                self.groq_client = Groq(api_key=groq_api_key)
                print("✅ Groq client initialized")
            except Exception as e:
                print(f"Warning: Could not initialize Groq: {e}")
    
    def extract_medical_entities(self, text: str) -> Dict:
        """Extract medical terms from text"""
        if not self.medical_ner or not text:
            return {"Disease": [], "Medication": [], "Procedure": []}
        
        try:
            entities = self.medical_ner(text[:1500])  # Limit text for speed
            medical_terms = {"Disease": [], "Medication": [], "Procedure": []}
            
            for entity in entities:
                if entity['score'] > 0.7:
                    category = entity['entity_group']
                    term = entity['word'].strip()
                    
                    if category in medical_terms and term not in medical_terms[category]:
                        medical_terms[category].append(term[:50])  # Limit term length
            
            return medical_terms
        except Exception as e:
            print(f"NER extraction error: {e}")
            return {"Disease": [], "Medication": [], "Procedure": []}
    
    def generate_ai_summary(self, text: str, medical_terms: Dict) -> str:
        """Generate patient-friendly summary using AI"""
        if not self.groq_client or not ENABLE_GROQ:
            return self.generate_rule_based_summary(text, medical_terms)
        
        try:
            # Prepare prompt
            terms_str = ""
            for category, terms in medical_terms.items():
                if terms:
                    terms_str += f"{category}: {', '.join(terms[:3])}\n"
            
            prompt = f"""Please summarize this medical report in simple, patient-friendly language:

Report: {text[:1000]}

Key medical terms found: {terms_str}

Provide a summary with:
1. Simple overview in 1-2 sentences
2. Key findings in plain language
3. What the patient should do next
Use bullet points and avoid medical jargon."""

            # Get AI response
            response = self.groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=AVAILABLE_GROQ_MODELS[0],
                max_tokens=500,
                temperature=0.3
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            print(f"AI summary error: {e}")
            return self.generate_rule_based_summary(text, medical_terms)
    
    def generate_rule_based_summary(self, text: str, medical_terms: Dict) -> str:
        """Generate summary without AI"""
        summary_parts = ["## 🏥 Medical Report Summary"]
        
        # Add key findings
        if any(medical_terms.values()):
            summary_parts.append("### 🔍 Key Findings:")
            for category, terms in medical_terms.items():
                if terms:
                    simplified = [self.simplify_term(term) for term in terms[:3]]
                    summary_parts.append(f"- **{category}**: {', '.join(simplified)}")
        
        # Add recommendations
        summary_parts.append("\n### 💡 Recommendations:")
        summary_parts.append("1. **Discuss** these findings with your healthcare provider")
        summary_parts.append("2. **Keep** this report for your medical records")
        summary_parts.append("3. **Follow up** as recommended by your doctor")
        
        # Disclaimer
        summary_parts.append("\n---")
        summary_parts.append("**Note**: This is an AI-generated summary for educational purposes only.")
        
        return "\n".join(summary_parts)
    
    def chat_about_report(self, question: str, report_text: str, medical_terms: Dict) -> str:
        """Answer questions about the medical report"""
        if not self.groq_client or not ENABLE_GROQ:
            return self._simple_chat_response(question, medical_terms)
        
        try:
            terms_str = "\n".join([f"{k}: {', '.join(v[:3])}" for k, v in medical_terms.items() if v])
            
            prompt = f"""You are a medical assistant. Answer this question based on the medical report.

Medical Report Context: {report_text[:800]}
Medical Terms Found: {terms_str}

Question: {question}

Provide a helpful, simple answer. If you cannot answer based on the report, suggest consulting a doctor."""

            response = self.groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=AVAILABLE_GROQ_MODELS[0],
                max_tokens=300,
                temperature=0.4
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            print(f"Chat error: {e}")
            return self._simple_chat_response(question, medical_terms)
    
    def _simple_chat_response(self, question: str, medical_terms: Dict) -> str:
        """Simple fallback chat response"""
        if "what does" in question.lower() or "mean" in question.lower():
            for term_list in medical_terms.values():
                for term in term_list:
                    if term.lower() in question.lower():
                        simplified = self.simplify_term(term)
                        return f"**{term}** typically means: {simplified}. For details specific to your case, please consult your doctor."
        
        return "I recommend discussing this with your healthcare provider for personalized advice."
    
    def simplify_term(self, term: str) -> str:
        """Simplify medical term"""
        return self.medical_glossary.get(term.lower(), term)

# ==================== MAIN APPLICATION ====================

class MedicalReportSummarizer:
    def __init__(self):
        self.document_processor = MultiFormatDocumentProcessor()
        self.ai_assistant = EnhancedMedicalAIAssistant(GROQ_API_KEY)
        self.consultant_search = PakistanConsultantSearch()
        self.current_report = None
    
    def process_document(self, file):
        """Process uploaded document"""
        if file is None:
            return self._get_placeholder_outputs()
        
        try:
            # Get file info
            file_path = file.name
            file_name = os.path.basename(file_path)
            file_ext = os.path.splitext(file_name)[1].lower().replace('.', '')
            
            if not file_ext:
                file_ext = 'txt'
            
            print(f"Processing: {file_name}")
            
            # Extract text
            extracted_text = self.document_processor.extract_text(file_path, file_ext)
            
            if extracted_text.startswith("Error") or len(extracted_text) < 10:
                return [
                    "❌ Could not extract meaningful text from the document. Please try a different file.",
                    "Unable to assess without text.",
                    "No medical terms found.",
                    "", []
                ]
            
            # Store current report
            self.current_report = extracted_text
            
            # Extract medical terms
            medical_terms = self.ai_assistant.extract_medical_entities(extracted_text)
            
            # Generate summary
            summary = self.ai_assistant.generate_ai_summary(extracted_text, medical_terms)
            
            # Analyze seriousness
            seriousness = self.ai_assistant.seriousness_analyzer.analyze_seriousness(extracted_text)
            seriousness_display = self._format_seriousness(seriousness)
            
            # Format medical terms
            terms_display = self._format_medical_terms(medical_terms)
            
            return [
                summary,
                seriousness_display,
                terms_display,
                "",  # Clear chat input
                []   # Clear chat history
            ]
            
        except Exception as e:
            print(f"Error: {e}")
            return [
                f"❌ Error processing document: {str(e)}",
                "Assessment failed",
                "Term extraction failed",
                "", []
            ]
    
    def handle_chat(self, question, chat_history):
        """Handle chat questions"""
        if not self.current_report:
            return "Please upload a medical report first.", chat_history
        
        # Get medical terms from current report
        medical_terms = self.ai_assistant.extract_medical_entities(self.current_report)
        
        # Get AI response
        response = self.ai_assistant.chat_about_report(
            question, 
            self.current_report, 
            medical_terms
        )
        
        # Update chat history
        chat_history.append((question, response))
        
        return "", chat_history
    
    def search_consultants(self, city, specialty, consultation_type):
        """Search for consultants"""
        online_only = (consultation_type == "Online Only")
        
        if not specialty:
            return "Please select a specialty."
        
        results = self.consultant_search.search_consultants(
            specialty, 
            city, 
            online_only
        )
        
        if not results:
            return "No consultants found. Try different search criteria."
        
        # Format results
        result_text = f"## Found {len(results)} Consultants\n\n"
        for doctor in results:
            result_text += self.consultant_search.format_doctor_card(doctor)
        
        return result_text
    
    def get_specialty_suggestions(self):
        """Get specialty suggestions based on current report"""
        if not self.current_report:
            return []
        
        medical_terms = self.ai_assistant.extract_medical_entities(self.current_report)
        
        # Flatten all terms
        all_terms = []
        for terms in medical_terms.values():
            all_terms.extend(terms)
        
        return self.consultant_search.suggest_specialties(all_terms)
    
    def _format_seriousness(self, seriousness_data: Dict) -> str:
        """Format seriousness assessment"""
        level = seriousness_data["level"]
        score = seriousness_data["score"]
        recommendation = seriousness_data["recommendation"]
        
        color_icon = {
            "High": "🔴",
            "Medium": "🟡", 
            "Low": "🟢"
        }.get(level, "⚪")
        
        return f"""
{color_icon} **Seriousness Level**: {level}

**Risk Score**: {score}/10

**Recommendation**:
{recommendation}
"""
    
    def _format_medical_terms(self, medical_terms: Dict) -> str:
        """Format medical terms for display"""
        if not any(medical_terms.values()):
            return "No specific medical terms identified."
        
        result = "### 🔍 Medical Terms Found:\n"
        for category, terms in medical_terms.items():
            if terms:
                result += f"\n**{category}:**\n"
                for term in terms[:5]:  # Limit to 5 terms per category
                    simplified = self.ai_assistant.simplify_term(term)
                    if simplified != term:
                        result += f"- {term} → *{simplified}*\n"
                    else:
                        result += f"- {term}\n"
        
        return result
    
    def _get_placeholder_outputs(self):
        return [
            "## 🏥 Medical Report Summarizer\n\nUpload a medical report to get started.",
            "Seriousness assessment will appear here.",
            "Medical terms will be extracted here.",
            "", []
        ]

# ==================== GRADIO INTERFACE ====================

def create_interface():
    """Create Gradio interface"""
    
    summarizer = MedicalReportSummarizer()
    
    with gr.Blocks(
        title="Medical Report Summarizer - Pakistan",
        theme=gr.themes.Soft(),
        css=".gradio-container {max-width: 1200px !important}"
    ) as demo:
        
        gr.Markdown("""
        # 🏥 AI Medical Report Summarizer - Pakistan
        **Upload medical reports, get AI summaries, and find Pakistani consultants**
        """)
        
        with gr.Tabs():
            # Tab 1: Medical Report Analysis
            with gr.Tab("📄 Analyze Medical Report"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("### 📤 Upload Medical Report")
                        file_input = gr.File(
                            label="Upload Document",
                            file_types=[
                                ".pdf", ".txt", ".docx", 
                                ".jpg", ".jpeg", ".png"
                            ],
                            type="filepath"
                        )
                        
                        gr.Markdown("""
                        ### 📋 Supported Formats:
                        - **PDF documents**
                        - **Text files** (.txt)
                        - **Word documents** (.docx)
                        - **Images** (.jpg, .png, etc.)
                        
                        *Note: Image files use OCR for text extraction*
                        """)
                    
                    with gr.Column(scale=2):
                        gr.Markdown("### 📝 AI Summary")
                        summary_output = gr.Markdown(
                            value="Upload a document to see the summary..."
                        )
                        
                        with gr.Accordion("🔍 Seriousness Assessment", open=True):
                            seriousness_output = gr.Markdown()
                        
                        with gr.Accordion("💊 Medical Terms Found", open=False):
                            terms_output = gr.Markdown()
                
                gr.Markdown("---")
                gr.Markdown("### 💬 Ask Questions About Your Report")
                
                chatbot = gr.Chatbot(
                    label="Medical Assistant Chat",
                    height=300,
                    bubble_full_width=False
                )
                
                with gr.Row():
                    chat_input = gr.Textbox(
                        label="Type your question...",
                        placeholder="Example: What does this finding mean?",
                        scale=4
                    )
                    send_btn = gr.Button("Send", variant="primary", scale=1)
            
            # Tab 2: Consultant Search
            with gr.Tab("🩺 Find Pakistani Consultants"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("### 🔍 Search Filters")
                        
                        # Auto-suggest from current report
                        gr.Markdown("**Based on your medical report:**")
                        auto_suggest_btn = gr.Button(
                            "🔄 Get Suggested Specialties",
                            variant="secondary"
                        )
                        suggested_specialties = gr.Dropdown(
                            label="Suggested Specialties",
                            choices=[],
                            interactive=True
                        )
                        
                        gr.Markdown("---")
                        gr.Markdown("**Manual Search:**")
                        
                        city_input = gr.Dropdown(
                            label="🏙️ City",
                            choices=["Any", "Online Only", "Karachi", "Lahore", "Islamabad"],
                            value="Any"
                        )
                        
                        specialty_input = gr.Dropdown(
                            label="🎯 Specialty",
                            choices=[
                                "General Physician", "Cardiology", "Endocrinology",
                                "Pediatrics", "Gynecology", "Dermatology"
                            ],
                            value="General Physician"
                        )
                        
                        consultation_type = gr.Radio(
                            label="💻 Consultation Type",
                            choices=["Any", "Online Only", "In-Person Only"],
                            value="Any"
                        )
                        
                        search_btn = gr.Button(
                            "🔍 Search Consultants", 
                            variant="primary",
                            size="lg"
                        )
                    
                    with gr.Column(scale=2):
                        gr.Markdown("### 👨‍⚕️ Consultant Results")
                        consultant_results = gr.Markdown(
                            value="Enter search criteria to find consultants..."
                        )
            
            # Tab 3: Healthcare Info
            with gr.Tab("ℹ️ Pakistan Healthcare Info"):
                gr.Markdown("""
                # 🇵🇰 Pakistan Healthcare Resources
                
                ## 🏥 Major Hospitals
                - **Aga Khan University Hospital** - Karachi
                - **Shaukat Khanum Memorial Hospital** - Lahore, Karachi
                - **Pakistan Institute of Medical Sciences (PIMS)** - Islamabad
                - **Civil Hospital** - Major cities
                
                ## 📞 Emergency Services
                - **Rescue 1122**: Nationwide emergency service
                - **Edhi Foundation**: 115-123-321
                - **Ambulance**: 1122 (in most areas)
                
                ## 💻 Telemedicine Services
                - **Sehat Kahani**: Online consultations
                - **Marham.pk**: Doctor appointments
                - **Oladdoctor**: Video consultations
                
                ## 💰 Typical Fees
                - General Physician: ₨ 1,000 - ₨ 2,000
                - Specialists: ₨ 2,000 - ₨ 4,000
                - Online: ₨ 1,000 - ₨ 2,500
                
                ## 🆘 Important Notes
                1. Always verify doctor credentials
                2. Keep medical records organized
                3. Ask about payment options
                4. For emergencies, go to nearest hospital
                """)
        
        # ===== EVENT HANDLERS =====
        
        # Process document
        file_input.change(
            fn=summarizer.process_document,
            inputs=[file_input],
            outputs=[
                summary_output,
                seriousness_output,
                terms_output,
                chat_input,
                chatbot
            ]
        )
        
        # Chat functionality
        def send_message(question, history):
            return summarizer.handle_chat(question, history)
        
        chat_input.submit(
            fn=send_message,
            inputs=[chat_input, chatbot],
            outputs=[chat_input, chatbot]
        )
        
        send_btn.click(
            fn=send_message,
            inputs=[chat_input, chatbot],
            outputs=[chat_input, chatbot]
        )
        
        # Consultant search
        search_btn.click(
            fn=summarizer.search_consultants,
            inputs=[city_input, specialty_input, consultation_type],
            outputs=[consultant_results]
        )
        
        # Auto-suggest specialties
        def update_suggestions():
            suggestions = summarizer.get_specialty_suggestions()
            if suggestions:
                return gr.update(choices=suggestions, value=suggestions[0])
            return gr.update(choices=[], value=None)
        
        auto_suggest_btn.click(
            fn=update_suggestions,
            outputs=[suggested_specialties]
        )
        
        # Use suggested specialty
        suggested_specialties.change(
            fn=lambda x: x,
            inputs=[suggested_specialties],
            outputs=[specialty_input]
        )
    
    return demo

# ==================== DEPLOYMENT SETUP ====================

def main():
    """Main entry point for deployment"""
    
    # Print configuration status
    print("=" * 50)
    print("Medical Report Summarizer - Deployment Status")
    print("=" * 50)
    
    print(f"📁 File Processing:")
    print(f"  PDF: {'✅' if PDF_AVAILABLE else '❌'}")
    print(f"  OCR: {'✅' if OCR_AVAILABLE else '❌'}")
    print(f"  DOCX: {'✅' if DOCX_AVAILABLE else '❌'}")
    
    print(f"\n🤖 AI Features:")
    print(f"  Transformers: {'✅' if TRANSFORMERS_AVAILABLE else '❌'}")
    print(f"  Groq AI: {'✅' if GROQ_AVAILABLE else '❌'}")
    print(f"  Groq Enabled: {'✅' if ENABLE_GROQ else '❌ (Set GROQ_API_KEY env var)'}")
    
    print(f"\n🩺 Medical Analysis:")
    print(f"  NER Model: {MEDICAL_NER_MODEL}")
    print(f"  Consultant DB: {len(PakistanConsultantSearch().doctors_database)} doctors")
    
    print("\n" + "=" * 50)
    print("🚀 Starting application...")
    
    # Create and launch interface
    demo = create_interface()
    
    # For Hugging Face Spaces
    if os.environ.get("SPACE_ID"):
        # Running on Hugging Face
        demo.launch(debug=False, share=False)
    else:
        # Running locally
        demo.launch(
            server_name="0.0.0.0",
            server_port=7860,
            share=False,
            debug=False
        )

# ==================== APP.PY FOR HUGGING FACE ====================

# Hugging Face Spaces requires the app to be accessible as `demo`
demo = create_interface()

# For local testing
if __name__ == "__main__":
    main()
